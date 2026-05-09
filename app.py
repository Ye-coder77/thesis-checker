import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import tempfile

st.set_page_config(page_title="论文格式检测工具（终极版）", layout="wide")
st.title("📄 学位论文格式检测工具（终极版）")

uploaded_file = st.file_uploader("上传论文（.docx）", type=["docx"])


# ===========================
# 工具函数（处理Word样式继承）
# ===========================
def get_font(run, para):
    return run.font.name or para.style.font.name

def get_size(run, para):
    if run.font.size:
        return run.font.size.pt
    if para.style.font.size:
        return para.style.font.size.pt
    return None

def get_line_spacing(para):
    pf = para.paragraph_format
    if pf.line_spacing_rule:
        return pf.line_spacing_rule
    if para.style and para.style.paragraph_format.line_spacing_rule:
        return para.style.paragraph_format.line_spacing_rule
    return None

def check_font(run, para, target):
    font = get_font(run, para)
    return font and target in font

def check_size(run, para, target):
    size = get_size(run, para)
    return size and abs(size - target) < 0.5

def get_spacing_before(para):
    if para.paragraph_format.space_before:
        return para.paragraph_format.space_before.pt
    if para.style and para.style.paragraph_format.space_before:
        return para.style.paragraph_format.space_before.pt
    return None

def get_spacing_after(para):
    if para.paragraph_format.space_after:
        return para.paragraph_format.space_after.pt
    if para.style and para.style.paragraph_format.space_after:
        return para.style.paragraph_format.space_after.pt
    return None

# ===========================
# 生成检测报告（Word）
# ===========================
def generate_report(results):
    doc = Document()
    doc.add_heading("论文格式检测报告", 0)

    for i, content in results.items():
        doc.add_paragraph(f"第{i}段：{content['text']}")
        for err in content["errors"]:
            doc.add_paragraph(f"- {err}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# ===========================
# 生成标注论文
# ===========================
def generate_marked_doc(original_doc, results):
    doc = original_doc

    for i, para in enumerate(doc.paragraphs, start=1):
        if i in results:
            # 高亮原文
            for run in para.runs:
                run.font.highlight_color = 7  # 黄色

            # 添加问题说明
            note = "\n【格式问题】\n"
            for err in results[i]["errors"]:
                note += f"- {err}\n"

            para.add_run(note)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# ===========================
# 分类函数（稳定版）
# ===========================
def classify(text):
    text = text.strip()

    # 摘要
    if text in ["摘要", "摘 要"]:
        return "cn_abstract_title"
    if text == "ABSTRACT":
        return "en_abstract_title"
    if text.startswith("关键词"):
        return "cn_keywords"
    if text.startswith("KEY WORDS"):
        return "en_keywords"

    # 标题（优先）
    if re.match(r'^\d+\.\d+\.\d+', text):
        return "title3"
    if re.match(r'^\d+\.\d+', text):
        return "title2"
    if re.match(r'^第\d+章', text):
        return "title1"
    if re.match(r'^\d+\.\s*', text):
        return "wrong_title1"
    if re.match(r'^（\d+）', text):
        return "title4"

    # 目录（放后面）
    if text == "目录":
        return "toc_title"
    if re.match(r'^\d+(\.\d+)*\s+.*\s+\d+$', text):
        return "toc_item"

    # 图表
    if text.startswith("图"):
        return "figure"
    if text.startswith("表"):
        return "table"

    # 正文
    if re.search(r'[a-zA-Z]', text):
        return "en_body"
    return "cn_body"


# ===========================
# 核心检测（完整版）
# ===========================
def check(doc):
    results = {}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        ptype = classify(text)
        pf = para.paragraph_format
        errors = []
        prev = doc.paragraphs[i-1].text.strip() if i > 0 else ""

        line_rule = get_line_spacing(para)

        # ===== 中文摘要标题 =====
        if ptype == "cn_abstract_title":
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append("摘要标题应居中")

            for run in para.runs:
                if not check_font(run, para, "黑体"):
                    errors.append("摘要标题应为黑体")
                if not check_size(run, para, 16):
                    errors.append("摘要标题应为三号（16pt）")

        # ===== 中文摘要正文 =====
        if prev in ["摘要", "摘 要"] and ptype == "cn_body":
            for run in para.runs:
                if not check_font(run, para, "宋体"):
                    errors.append("摘要正文应为宋体")
                if not check_size(run, para, 12):
                    errors.append("摘要正文应为小四（12pt）")

            if line_rule != 1:
                errors.append("摘要正文应为1.5倍行距")

            if pf.first_line_indent is None:
                errors.append("摘要正文应首行缩进2字符")

        # ===== 英文摘要标题 =====
        if ptype == "en_abstract_title":
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append("ABSTRACT应居中")

            if not any(run.bold for run in para.runs):
                errors.append("ABSTRACT应加粗")

            for run in para.runs:
                if not check_font(run, para, "Times"):
                    errors.append("ABSTRACT应为Times New Roman")
                if not check_size(run, para, 16):
                    errors.append("ABSTRACT应为三号（16pt）")

        # ===== 英文摘要正文 =====
        if prev == "ABSTRACT" and ptype == "en_body":
            if line_rule != 2:
                errors.append("英文摘要应为2倍行距")

        # ===== 中文正文 =====
        if ptype == "cn_body":
            for run in para.runs:
                if not check_font(run, para, "宋体"):
                    errors.append("中文正文应为宋体")
                if not check_size(run, para, 12):
                    errors.append("中文正文应为小四（12pt）")

            if line_rule != 1:
                errors.append("正文应为1.5倍行距")

            if pf.first_line_indent is None:
                errors.append("正文应首行缩进2字符")

        # ===== 英文正文 =====
        if ptype == "en_body":
            for run in para.runs:
                if not check_font(run, para, "Times"):
                    errors.append("英文正文应为Times New Roman")

        # ===== 一级标题 =====
        if ptype == "title1":
            if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append("一级标题必须居中")

            before = get_spacing_before(para)
            after = get_spacing_after(para)

            # 1行 ≈ 14–18pt
            if before is None or before < 12:
                errors.append("一级标题段前应空1行")

            if after is None or after < 12:
                errors.append("一级标题段后应空1行")

        # ===== 错误一级标题 =====
        if ptype == "wrong_title1":
            errors.append("一级标题必须为“第X章”")

        # ===== 二级标题 =====
        if ptype == "title2":
            before = get_spacing_before(para)
            after = get_spacing_after(para)

            # 0.5行 ≈ 6pt
            if before is None or before < 5:
                errors.append("二级标题段前应为0.5行")

            if after is None or after < 5:
                errors.append("二级标题段后应为0.5行")

        # ===== 图 =====
        if ptype == "figure":
            if not re.match(r'^图\d+[-\.]\d+', text):
                errors.append("图编号错误")

        # ===== 表 =====
        if ptype == "table":
            if not re.match(r'^表\d+\.\d+', text):
                errors.append("表编号错误")

        # ===== 目录 =====
        if ptype == "toc_item":
            if text.count(".") > 2:
                errors.append("目录最多三级标题")

        if errors:
            results[i + 1] = {
                "text": text,
                "type": ptype,
                "errors": list(set(errors))
            }

    return results


# ===========================
# UI
# ===========================
if uploaded_file:
    doc = Document(uploaded_file)
    results = check(doc)

    categories = {
        "cn_abstract_title": "摘要",
        "cn_keywords": "摘要",
        "en_abstract_title": "摘要",
        "en_keywords": "摘要",
        "cn_body": "正文",
        "en_body": "正文",
        "title1": "标题",
        "title2": "标题",
        "title3": "标题",
        "title4": "标题",
        "wrong_title1": "标题",
        "figure": "图表",
        "table": "图表",
        "toc_title": "目录",
        "toc_item": "目录"
    }

    grouped = {"摘要": [], "正文": [], "标题": [], "图表": [], "目录": []}

    # 分类汇总
for para, content in results.items():
    group = categories.get(content["type"], "正文")
    grouped[group].append((para, content))

st.error(f"⚠️ 共发现 {len(results)} 处问题")

# ===========================
# 展示检测结果
# ===========================
for group_name, items in grouped.items():
    if items:
        with st.expander(f"📌 {group_name}部分（{len(items)}项）"):
            for para, content in items:

                # 段落标题
                st.markdown(f"**第{para}段：{content['text']}**")

                # 错误列表
                for err in content["errors"]:
                    st.write(f"👉 {err}")

                # 分隔（让UI清晰）
                st.markdown("---")

# ===========================
# ⬇️ 新增功能（导出 + 标注）
# ===========================
st.divider()

# 下载检测报告
report_path = generate_report(results)
with open(report_path, "rb") as f:
    st.download_button(
        "📥 下载检测报告",
        f,
        file_name="论文格式检测报告.docx"
    )

# 标注论文
if st.button("🖍 生成标注论文"):
    doc2 = Document(uploaded_file)
    marked_path = generate_marked_doc(doc2, results)

    with open(marked_path, "rb") as f:
        st.download_button(
            "📥 下载标注论文",
            f,
            file_name="标注论文.docx"
        )
